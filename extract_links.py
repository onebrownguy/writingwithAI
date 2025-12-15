
import zipfile
import re
import sys
import xml.etree.ElementTree as ET

docx_path = r"Writing with AI Detailed Schedule Fall 2025 (1).docx"

import docx

def extract_content(file_path):
    try:
        doc = docx.Document(file_path)
        print("--- CONTENT START ---")
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                print(f"[PARA] {para.text.strip()}")
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            print(f"[TABLE] {para.text.strip()}")
        print("--- CONTENT END ---")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_content("Writing with AI Detailed Schedule Fall 2025 (1).docx")
